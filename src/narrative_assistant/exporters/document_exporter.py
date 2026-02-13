"""
Exportador de documentos completos a DOCX y PDF.

Genera documentos profesionales con:
- Portada con titulo del proyecto
- Indice automatico
- Fichas de personajes
- Alertas/errores encontrados
- Timeline narrativo
- Grafo de relaciones (como imagen)
- Guia de estilo

Formatos:
- DOCX: Usando python-docx (ya instalado)
- PDF: Usando reportlab o weasyprint
"""

import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..core.errors import ErrorSeverity, NarrativeError
from ..core.result import Result

logger = logging.getLogger(__name__)


class ExportSection(Enum):
    """Secciones disponibles para exportar."""

    COVER = "cover"
    TOC = "toc"
    CHARACTER_SHEETS = "character_sheets"
    ALERTS = "alerts"
    TIMELINE = "timeline"
    RELATIONSHIPS = "relationships"
    STYLE_GUIDE = "style_guide"
    STATISTICS = "statistics"


@dataclass
class ExportOptions:
    """Opciones de exportacion de documento."""

    # Secciones a incluir
    include_cover: bool = True
    include_toc: bool = True
    include_character_sheets: bool = True
    include_alerts: bool = True
    include_timeline: bool = True
    include_relationships: bool = True
    include_style_guide: bool = True
    include_statistics: bool = True

    # Filtros
    only_main_characters: bool = True
    only_open_alerts: bool = True
    min_alert_severity: str = "info"  # info, warning, error, critical

    # Formato
    include_excerpts: bool = True
    max_excerpts_per_alert: int = 3

    def get_enabled_sections(self) -> list[ExportSection]:
        """Devuelve las secciones habilitadas."""
        sections = []
        if self.include_cover:
            sections.append(ExportSection.COVER)
        if self.include_toc:
            sections.append(ExportSection.TOC)
        if self.include_statistics:
            sections.append(ExportSection.STATISTICS)
        if self.include_character_sheets:
            sections.append(ExportSection.CHARACTER_SHEETS)
        if self.include_alerts:
            sections.append(ExportSection.ALERTS)
        if self.include_timeline:
            sections.append(ExportSection.TIMELINE)
        if self.include_relationships:
            sections.append(ExportSection.RELATIONSHIPS)
        if self.include_style_guide:
            sections.append(ExportSection.STYLE_GUIDE)
        return sections


@dataclass
class ProjectExportData:
    """Datos del proyecto para exportar."""

    project_id: int
    project_name: str
    description: str = ""

    # Estadisticas
    word_count: int = 0
    chapter_count: int = 0
    entity_count: int = 0
    alert_count: int = 0

    # Datos para secciones
    characters: list[dict] = field(default_factory=list)
    alerts: list[dict] = field(default_factory=list)
    timeline_events: list[dict] = field(default_factory=list)
    relationships: list[dict] = field(default_factory=list)
    style_guide: dict | None = None

    # Metadata
    generated_at: datetime = field(default_factory=datetime.now)


class DocumentExporter:
    """
    Exportador de documentos completos a DOCX y PDF.

    Uso:
        exporter = DocumentExporter()
        docx_bytes = exporter.export_to_docx(project_data, options)
        pdf_bytes = exporter.export_to_pdf(project_data, options)
    """

    def __init__(self):
        """Inicializa el exportador."""
        self._setup_styles_cache = {}

    def export_to_docx(
        self,
        data: ProjectExportData,
        options: ExportOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta el proyecto a formato DOCX.

        Args:
            data: Datos del proyecto a exportar
            options: Opciones de exportacion

        Returns:
            Result con bytes del documento DOCX
        """
        if options is None:
            options = ExportOptions()

        try:
            # Crear documento
            doc = Document()
            self._setup_styles(doc)

            # Generar secciones
            sections = options.get_enabled_sections()

            for section in sections:
                if section == ExportSection.COVER:
                    self._add_cover_page(doc, data)
                elif section == ExportSection.TOC:
                    self._add_table_of_contents(doc)
                elif section == ExportSection.STATISTICS:
                    self._add_statistics_section(doc, data)
                elif section == ExportSection.CHARACTER_SHEETS:
                    self._add_character_sheets_section(doc, data, options)
                elif section == ExportSection.ALERTS:
                    self._add_alerts_section(doc, data, options)
                elif section == ExportSection.TIMELINE:
                    self._add_timeline_section(doc, data)
                elif section == ExportSection.RELATIONSHIPS:
                    self._add_relationships_section(doc, data)
                elif section == ExportSection.STYLE_GUIDE:
                    self._add_style_guide_section(doc, data)

            # Footer
            self._add_footer(doc, data)

            # Guardar a bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(f"Exported DOCX for project {data.project_id}")
            return Result.success(buffer.getvalue())

        except Exception as e:
            error = NarrativeError(
                message=f"Failed to export DOCX: {str(e)}",
                severity=ErrorSeverity.ERROR,
                user_message=f"Error exportando a DOCX: {str(e)}",
            )
            logger.error(f"Error exporting DOCX: {e}", exc_info=True)
            return Result.failure(error)

    def export_to_pdf(
        self,
        data: ProjectExportData,
        options: ExportOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta el proyecto a formato PDF.

        Usa reportlab para generacion directa de PDF.

        Args:
            data: Datos del proyecto a exportar
            options: Opciones de exportacion

        Returns:
            Result con bytes del documento PDF
        """
        if options is None:
            options = ExportOptions()

        try:
            # Intentar importar reportlab
            from reportlab.lib.colors import HexColor, black, gray
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                ListFlowable,
                ListItem,
                PageBreak,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )

            buffer = io.BytesIO()

            # Crear documento PDF
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=2 * cm,
                leftMargin=2 * cm,
                topMargin=2 * cm,
                bottomMargin=2 * cm,
            )

            # Estilos
            styles = getSampleStyleSheet()

            # Estilo titulo
            styles.add(
                ParagraphStyle(
                    name="CoverTitle",
                    parent=styles["Title"],
                    fontSize=28,
                    spaceAfter=30,
                    alignment=TA_CENTER,
                    textColor=HexColor("#1a1a2e"),
                )
            )

            # Estilo subtitulo
            styles.add(
                ParagraphStyle(
                    name="CoverSubtitle",
                    parent=styles["Normal"],
                    fontSize=14,
                    spaceAfter=20,
                    alignment=TA_CENTER,
                    textColor=HexColor("#4a4a6a"),
                )
            )

            # Estilo Heading1
            styles.add(
                ParagraphStyle(
                    name="CustomHeading1",
                    parent=styles["Heading1"],
                    fontSize=18,
                    spaceBefore=20,
                    spaceAfter=12,
                    textColor=HexColor("#1a1a2e"),
                )
            )

            # Estilo Heading2
            styles.add(
                ParagraphStyle(
                    name="CustomHeading2",
                    parent=styles["Heading2"],
                    fontSize=14,
                    spaceBefore=15,
                    spaceAfter=8,
                    textColor=HexColor("#2d2d4a"),
                )
            )

            # Estilo para alertas
            styles.add(
                ParagraphStyle(
                    name="AlertCritical",
                    parent=styles["Normal"],
                    fontSize=10,
                    leftIndent=20,
                    textColor=HexColor("#dc3545"),
                )
            )

            styles.add(
                ParagraphStyle(
                    name="AlertWarning",
                    parent=styles["Normal"],
                    fontSize=10,
                    leftIndent=20,
                    textColor=HexColor("#ffc107"),
                )
            )

            story = []
            sections = options.get_enabled_sections()

            # Portada
            if ExportSection.COVER in sections:
                story.extend(self._build_pdf_cover(data, styles))
                story.append(PageBreak())

            # Estadisticas
            if ExportSection.STATISTICS in sections:
                story.extend(self._build_pdf_statistics(data, styles))
                story.append(Spacer(1, 20))

            # Fichas de personajes
            if ExportSection.CHARACTER_SHEETS in sections:
                story.extend(self._build_pdf_characters(data, options, styles))
                story.append(Spacer(1, 20))

            # Alertas
            if ExportSection.ALERTS in sections:
                story.extend(self._build_pdf_alerts(data, options, styles))
                story.append(Spacer(1, 20))

            # Timeline
            if ExportSection.TIMELINE in sections:
                story.extend(self._build_pdf_timeline(data, styles))
                story.append(Spacer(1, 20))

            # Relaciones
            if ExportSection.RELATIONSHIPS in sections:
                story.extend(self._build_pdf_relationships(data, styles))
                story.append(Spacer(1, 20))

            # Guia de estilo
            if ExportSection.STYLE_GUIDE in sections:
                story.extend(self._build_pdf_style_guide(data, styles))

            # Footer
            story.append(Spacer(1, 30))
            story.append(
                Paragraph(
                    f"Generado por Narrative Assistant el {data.generated_at.strftime('%Y-%m-%d %H:%M')}",
                    styles["Normal"],
                )
            )

            # Construir PDF
            doc.build(story)
            buffer.seek(0)

            logger.info(f"Exported PDF for project {data.project_id}")
            return Result.success(buffer.getvalue())

        except ImportError:
            error = NarrativeError(
                message="reportlab not installed",
                severity=ErrorSeverity.ERROR,
                user_message="Para exportar a PDF es necesario instalar reportlab: pip install reportlab",
            )
            return Result.failure(error)
        except Exception as e:
            error = NarrativeError(
                message=f"Failed to export PDF: {str(e)}",
                severity=ErrorSeverity.ERROR,
                user_message=f"Error exportando a PDF: {str(e)}",
            )
            logger.error(f"Error exporting PDF: {e}", exc_info=True)
            return Result.failure(error)

    # =========================================================================
    # DOCX Generation Methods
    # =========================================================================

    def _setup_styles(self, doc: Document) -> None:
        """Configura estilos del documento."""
        styles = doc.styles

        # Estilo para titulos de seccion
        try:
            title_style = styles.add_style("SectionTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(26, 26, 46)
            title_style.paragraph_format.space_before = Pt(24)
            title_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            pass  # Estilo ya existe

        # Estilo para subtitulos
        try:
            subtitle_style = styles.add_style("SubsectionTitle", WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.size = Pt(14)
            subtitle_style.font.bold = True
            subtitle_style.font.color.rgb = RGBColor(45, 45, 74)
            subtitle_style.paragraph_format.space_before = Pt(18)
            subtitle_style.paragraph_format.space_after = Pt(8)
        except ValueError:
            pass

        # Estilo para extractos
        try:
            excerpt_style = styles.add_style("Excerpt", WD_STYLE_TYPE.PARAGRAPH)
            excerpt_style.font.size = Pt(10)
            excerpt_style.font.italic = True
            excerpt_style.font.color.rgb = RGBColor(100, 100, 100)
            excerpt_style.paragraph_format.left_indent = Inches(0.5)
        except ValueError:
            pass

        # Estilo para alertas criticas
        try:
            critical_style = styles.add_style("AlertCritical", WD_STYLE_TYPE.PARAGRAPH)
            critical_style.font.size = Pt(11)
            critical_style.font.color.rgb = RGBColor(220, 53, 69)
        except ValueError:
            pass

        # Estilo para advertencias
        try:
            warning_style = styles.add_style("AlertWarning", WD_STYLE_TYPE.PARAGRAPH)
            warning_style.font.size = Pt(11)
            warning_style.font.color.rgb = RGBColor(255, 193, 7)
        except ValueError:
            pass

    def _add_cover_page(self, doc: Document, data: ProjectExportData) -> None:
        """Añade portada al documento."""
        # Espaciado superior
        for _ in range(3):
            doc.add_paragraph()

        # Titulo
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(data.project_name)
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(26, 26, 46)

        # Subtitulo
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Informe de Analisis Narrativo")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(74, 74, 106)

        # Descripcion
        if data.description:
            doc.add_paragraph()
            desc = doc.add_paragraph()
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = desc.add_run(data.description)
            run.font.size = Pt(12)
            run.font.italic = True

        # Espaciado
        for _ in range(5):
            doc.add_paragraph()

        # Fecha
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generado el {data.generated_at.strftime('%d de %B de %Y')}")
        run.font.size = Pt(12)

        # Salto de pagina
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Añade tabla de contenidos."""
        # Titulo
        toc_title = doc.add_paragraph()
        run = toc_title.add_run("Indice")
        run.bold = True
        run.font.size = Pt(18)

        # Nota sobre TOC
        note = doc.add_paragraph()
        run = note.add_run(
            "Nota: Para actualizar el indice en Word, haga clic derecho "
            "sobre el mismo y seleccione 'Actualizar campo'."
        )
        run.font.size = Pt(10)
        run.font.italic = True

        # Campo TOC (requiere actualizacion manual en Word)
        toc_para = doc.add_paragraph()
        # Añadir campo TOC usando XML
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = toc_para.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        run2 = toc_para.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run2._r.append(instr_text)

        run3 = toc_para.add_run()
        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_char_separate)

        run4 = toc_para.add_run("Actualice este campo para ver el indice")
        run4.font.italic = True

        run5 = toc_para.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fld_char_end)

        doc.add_page_break()

    def _add_statistics_section(self, doc: Document, data: ProjectExportData) -> None:
        """Añade seccion de estadisticas."""
        doc.add_heading("Estadisticas del Proyecto", level=1)

        # Tabla de estadisticas
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        stats = [
            ("Total de palabras", f"{data.word_count:,}"),
            ("Capitulos", str(data.chapter_count)),
            ("Entidades detectadas", str(data.entity_count)),
            ("Alertas totales", str(data.alert_count)),
            ("Fecha de generacion", data.generated_at.strftime("%Y-%m-%d %H:%M")),
        ]

        for i, (label, value) in enumerate(stats):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value

        doc.add_paragraph()

    def _add_character_sheets_section(
        self,
        doc: Document,
        data: ProjectExportData,
        options: ExportOptions,
    ) -> None:
        """Añade seccion de fichas de personajes."""
        doc.add_heading("Fichas de Personajes", level=1)

        if not data.characters:
            doc.add_paragraph("No se detectaron personajes en el manuscrito.")
            return

        # Filtrar personajes si es necesario
        characters = data.characters
        if options.only_main_characters:
            characters = [
                c for c in characters if c.get("importance") in ("principal", "critical", "high")
            ]

        if not characters:
            doc.add_paragraph("No hay personajes principales para mostrar.")
            return

        for char in characters:
            # Nombre del personaje
            doc.add_heading(char.get("canonical_name", "Sin nombre"), level=2)

            # Informacion basica
            if char.get("aliases"):
                p = doc.add_paragraph()
                p.add_run("Tambien conocido como: ").bold = True
                p.add_run(", ".join(char["aliases"]))

            p = doc.add_paragraph()
            p.add_run("Tipo: ").bold = True
            p.add_run(char.get("entity_type", "Personaje"))

            p = doc.add_paragraph()
            p.add_run("Importancia: ").bold = True
            p.add_run(char.get("importance", "Desconocida"))

            # Apariciones
            if char.get("mentions"):
                p = doc.add_paragraph()
                p.add_run("Menciones totales: ").bold = True
                p.add_run(str(char["mentions"].get("total_mentions", 0)))

                if char["mentions"].get("chapters"):
                    p = doc.add_paragraph()
                    p.add_run("Aparece en capitulos: ").bold = True
                    p.add_run(", ".join(str(c) for c in char["mentions"]["chapters"]))

            # Atributos fisicos
            if char.get("physical_attributes"):
                doc.add_heading("Atributos Fisicos", level=3)
                for attr in char["physical_attributes"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{attr.get('key', '').replace('_', ' ').title()}: ").bold = True
                    p.add_run(attr.get("value", ""))
                    if attr.get("confidence"):
                        p.add_run(f" (confianza: {attr['confidence']:.0%})")

            # Atributos psicologicos
            if char.get("psychological_attributes"):
                doc.add_heading("Atributos Psicologicos", level=3)
                for attr in char["psychological_attributes"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{attr.get('key', '').replace('_', ' ').title()}: ").bold = True
                    p.add_run(attr.get("value", ""))

            # Perfil de voz
            if char.get("voice_profile"):
                vp = char["voice_profile"]
                doc.add_heading("Perfil de Voz", level=3)

                if vp.get("total_interventions"):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("Intervenciones: ").bold = True
                    p.add_run(str(vp["total_interventions"]))

                if vp.get("formality_score") is not None:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("Formalidad: ").bold = True
                    p.add_run(f"{vp['formality_score']:.0%}")

                if vp.get("predominant_register"):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run("Registro: ").bold = True
                    p.add_run(vp["predominant_register"])

            doc.add_paragraph()  # Espaciado entre personajes

    def _add_alerts_section(
        self,
        doc: Document,
        data: ProjectExportData,
        options: ExportOptions,
    ) -> None:
        """Añade seccion de alertas."""
        doc.add_heading("Alertas e Inconsistencias", level=1)

        if not data.alerts:
            doc.add_paragraph("No se detectaron alertas en el manuscrito.")
            return

        # Filtrar alertas
        alerts = data.alerts
        if options.only_open_alerts:
            alerts = [a for a in alerts if a.get("status") in ("open", "active", "pending")]

        # Filtrar por severidad
        severity_order = {"info": 0, "warning": 1, "error": 2, "critical": 3}
        min_severity = severity_order.get(options.min_alert_severity, 0)
        alerts = [
            a for a in alerts if severity_order.get(a.get("severity", "info"), 0) >= min_severity
        ]

        if not alerts:
            doc.add_paragraph("No hay alertas que cumplan los criterios de filtrado.")
            return

        # Agrupar por severidad
        critical = [a for a in alerts if a.get("severity") == "critical"]
        errors = [a for a in alerts if a.get("severity") == "error"]
        warnings = [a for a in alerts if a.get("severity") == "warning"]
        info = [a for a in alerts if a.get("severity") == "info"]

        # Criticas
        if critical:
            doc.add_heading(f"Criticas ({len(critical)})", level=2)
            for alert in critical:
                self._add_alert_item(doc, alert, options, is_critical=True)

        # Errores
        if errors:
            doc.add_heading(f"Errores ({len(errors)})", level=2)
            for alert in errors:
                self._add_alert_item(doc, alert, options, is_error=True)

        # Advertencias
        if warnings:
            doc.add_heading(f"Advertencias ({len(warnings)})", level=2)
            for alert in warnings:
                self._add_alert_item(doc, alert, options)

        # Informativas
        if info:
            doc.add_heading(f"Informativas ({len(info)})", level=2)
            for alert in info:
                self._add_alert_item(doc, alert, options)

    def _add_alert_item(
        self,
        doc: Document,
        alert: dict,
        options: ExportOptions,
        is_critical: bool = False,
        is_error: bool = False,
    ) -> None:
        """Añade un item de alerta al documento."""
        # Titulo de la alerta
        p = doc.add_paragraph()
        title_run = p.add_run(alert.get("title", "Alerta sin titulo"))
        title_run.bold = True
        if is_critical:
            title_run.font.color.rgb = RGBColor(220, 53, 69)
        elif is_error:
            title_run.font.color.rgb = RGBColor(255, 100, 100)

        # Descripcion
        if alert.get("description"):
            p = doc.add_paragraph()
            p.add_run("Descripcion: ").bold = True
            p.add_run(alert["description"])

        # Explicacion
        if alert.get("explanation"):
            p = doc.add_paragraph()
            p.add_run("Explicacion: ").bold = True
            p.add_run(alert["explanation"])

        # Sugerencia
        if alert.get("suggestion"):
            p = doc.add_paragraph()
            p.add_run("Sugerencia: ").bold = True
            p.add_run(alert["suggestion"])

        # Ubicacion
        if alert.get("chapter"):
            p = doc.add_paragraph()
            p.add_run("Ubicacion: ").bold = True
            location = f"Capitulo {alert['chapter']}"
            if alert.get("scene"):
                location += f", Escena {alert['scene']}"
            p.add_run(location)

        # Extracto
        if options.include_excerpts and alert.get("excerpt"):
            p = doc.add_paragraph()
            p.add_run("Extracto: ").bold = True
            excerpt = alert["excerpt"][:200]
            if len(alert["excerpt"]) > 200:
                excerpt += "..."
            excerpt_para = doc.add_paragraph(style="Quote")
            excerpt_para.add_run(excerpt).italic = True

        # Confianza
        if alert.get("confidence"):
            p = doc.add_paragraph()
            p.add_run(f"Confianza: {alert['confidence']:.0%}")

        doc.add_paragraph()  # Espaciado

    def _add_timeline_section(self, doc: Document, data: ProjectExportData) -> None:
        """Añade seccion de timeline."""
        doc.add_heading("Linea Temporal", level=1)

        if not data.timeline_events:
            doc.add_paragraph("No se detectaron eventos temporales en el manuscrito.")
            return

        # Agrupar por capitulo
        events_by_chapter: dict[int, list] = {}
        for event in data.timeline_events:
            chapter = event.get("chapter", 0)
            if chapter not in events_by_chapter:
                events_by_chapter[chapter] = []
            events_by_chapter[chapter].append(event)

        for chapter in sorted(events_by_chapter.keys()):
            doc.add_heading(f"Capitulo {chapter}", level=2)

            for event in events_by_chapter[chapter]:
                p = doc.add_paragraph(style="List Bullet")

                # Fecha si existe
                if event.get("story_date"):
                    p.add_run(f"[{event['story_date']}] ").bold = True

                # Descripcion
                p.add_run(event.get("description", "Evento sin descripcion"))

                # Tipo narrativo
                narrative_order = event.get("narrative_order", "chronological")
                if narrative_order == "analepsis":
                    p.add_run(" [FLASHBACK]").italic = True
                elif narrative_order == "prolepsis":
                    p.add_run(" [FLASHFORWARD]").italic = True

    def _add_relationships_section(self, doc: Document, data: ProjectExportData) -> None:
        """Añade seccion de relaciones."""
        doc.add_heading("Relaciones entre Personajes", level=1)

        if not data.relationships:
            doc.add_paragraph("No se detectaron relaciones en el manuscrito.")
            return

        # Crear tabla de relaciones
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Encabezados
        headers = ["Personaje 1", "Relacion", "Personaje 2", "Intensidad"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Datos
        for rel in data.relationships:
            row = table.add_row()
            row.cells[0].text = rel.get("source_entity_name", "")
            row.cells[1].text = rel.get("relation_type", "").replace("_", " ").title()
            row.cells[2].text = rel.get("target_entity_name", "")
            intensity = rel.get("intensity", 0.5)
            row.cells[3].text = f"{intensity:.0%}"

    def _add_style_guide_section(self, doc: Document, data: ProjectExportData) -> None:
        """Añade seccion de guia de estilo."""
        doc.add_heading("Guia de Estilo", level=1)

        if not data.style_guide:
            doc.add_paragraph("No se genero guia de estilo para este proyecto.")
            return

        sg = data.style_guide

        # Decisiones de grafia
        if sg.get("spelling_decisions"):
            doc.add_heading("Decisiones de Grafia", level=2)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            headers = ["Forma Canonica", "Variantes", "Recomendacion"]
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            for decision in sg["spelling_decisions"]:
                row = table.add_row()
                row.cells[0].text = decision.get("canonical_form", "")
                row.cells[1].text = ", ".join(decision.get("variants", []))
                row.cells[2].text = decision.get("recommendation", "")

        # Personajes
        if sg.get("characters"):
            doc.add_heading("Personajes Registrados", level=2)
            for char in sg["characters"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(char.get("canonical_name", "")).bold = True
                if char.get("aliases"):
                    p.add_run(f" (alias: {', '.join(char['aliases'])})")

        # Analisis estilistico
        if sg.get("style_analysis"):
            sa = sg["style_analysis"]
            doc.add_heading("Analisis Estilistico", level=2)

            if sa.get("dialogue_style"):
                p = doc.add_paragraph()
                p.add_run("Estilo de dialogos: ").bold = True
                p.add_run(sa["dialogue_style"])

            if sa.get("number_style"):
                p = doc.add_paragraph()
                p.add_run("Estilo de numeros: ").bold = True
                p.add_run(sa["number_style"])

            # Recomendaciones
            if sa.get("recommendations"):
                doc.add_heading("Recomendaciones", level=3)
                for rec in sa["recommendations"]:
                    doc.add_paragraph(rec, style="List Bullet")

    def _add_footer(self, doc: Document, data: ProjectExportData) -> None:
        """Añade footer al documento."""
        doc.add_paragraph()
        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Documento generado por Narrative Assistant el "
            f"{data.generated_at.strftime('%Y-%m-%d %H:%M')}"
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    # =========================================================================
    # PDF Generation Methods
    # =========================================================================

    def _build_pdf_cover(self, data: ProjectExportData, styles) -> list:
        """Construye la portada para PDF."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Spacer(1, 100))
        story.append(Paragraph(data.project_name, styles["CoverTitle"]))
        story.append(Paragraph("Informe de Analisis Narrativo", styles["CoverSubtitle"]))

        if data.description:
            story.append(Spacer(1, 20))
            story.append(Paragraph(data.description, styles["Normal"]))

        story.append(Spacer(1, 100))
        story.append(
            Paragraph(
                f"Generado el {data.generated_at.strftime('%d de %B de %Y')}",
                styles["CoverSubtitle"],
            )
        )

        return story

    def _build_pdf_statistics(self, data: ProjectExportData, styles) -> list:
        """Construye seccion de estadisticas para PDF."""
        from reportlab.lib.colors import black, lightgrey
        from reportlab.platypus import Paragraph, Table, TableStyle

        story = []
        story.append(Paragraph("Estadisticas del Proyecto", styles["CustomHeading1"]))

        table_data = [
            ["Metrica", "Valor"],
            ["Total de palabras", f"{data.word_count:,}"],
            ["Capitulos", str(data.chapter_count)],
            ["Entidades detectadas", str(data.entity_count)],
            ["Alertas totales", str(data.alert_count)],
        ]

        table = Table(table_data, colWidths=[200, 150])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                    ("GRID", (0, 0), (-1, -1), 1, black),
                ]
            )
        )

        story.append(table)
        return story

    def _build_pdf_characters(
        self,
        data: ProjectExportData,
        options: ExportOptions,
        styles,
    ) -> list:
        """Construye seccion de personajes para PDF."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Paragraph("Fichas de Personajes", styles["CustomHeading1"]))

        if not data.characters:
            story.append(Paragraph("No se detectaron personajes.", styles["Normal"]))
            return story

        characters = data.characters
        if options.only_main_characters:
            characters = [
                c for c in characters if c.get("importance") in ("principal", "critical", "high")
            ]

        for char in characters:
            story.append(
                Paragraph(char.get("canonical_name", "Sin nombre"), styles["CustomHeading2"])
            )

            info_items = []
            if char.get("entity_type"):
                info_items.append(f"<b>Tipo:</b> {char['entity_type']}")
            if char.get("importance"):
                info_items.append(f"<b>Importancia:</b> {char['importance']}")
            if char.get("aliases"):
                info_items.append(f"<b>Alias:</b> {', '.join(char['aliases'])}")

            for item in info_items:
                story.append(Paragraph(item, styles["Normal"]))

            story.append(Spacer(1, 10))

        return story

    def _build_pdf_alerts(
        self,
        data: ProjectExportData,
        options: ExportOptions,
        styles,
    ) -> list:
        """Construye seccion de alertas para PDF."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Paragraph("Alertas e Inconsistencias", styles["CustomHeading1"]))

        if not data.alerts:
            story.append(Paragraph("No se detectaron alertas.", styles["Normal"]))
            return story

        alerts = data.alerts
        if options.only_open_alerts:
            alerts = [a for a in alerts if a.get("status") in ("open", "active", "pending")]

        severity_order = {"info": 0, "warning": 1, "error": 2, "critical": 3}
        min_severity = severity_order.get(options.min_alert_severity, 0)
        alerts = [
            a for a in alerts if severity_order.get(a.get("severity", "info"), 0) >= min_severity
        ]

        for alert in alerts[:20]:  # Limitar a 20 alertas
            severity = alert.get("severity", "info")
            title = alert.get("title", "Alerta")

            if severity == "critical":
                style = styles["AlertCritical"]
            elif severity == "warning":
                style = styles["AlertWarning"]
            else:
                style = styles["Normal"]

            story.append(Paragraph(f"<b>{title}</b>", style))

            if alert.get("description"):
                story.append(Paragraph(alert["description"], styles["Normal"]))

            story.append(Spacer(1, 8))

        return story

    def _build_pdf_timeline(self, data: ProjectExportData, styles) -> list:
        """Construye seccion de timeline para PDF."""
        from reportlab.platypus import Paragraph

        story = []
        story.append(Paragraph("Linea Temporal", styles["CustomHeading1"]))

        if not data.timeline_events:
            story.append(Paragraph("No se detectaron eventos temporales.", styles["Normal"]))
            return story

        for event in data.timeline_events[:30]:  # Limitar
            date_str = event.get("story_date", "")
            desc = event.get("description", "Evento")

            text = f"<b>{date_str}</b> - {desc}" if date_str else desc
            story.append(Paragraph(text, styles["Normal"]))

        return story

    def _build_pdf_relationships(self, data: ProjectExportData, styles) -> list:
        """Construye seccion de relaciones para PDF."""
        from reportlab.lib.colors import black, lightgrey
        from reportlab.platypus import Paragraph, Table, TableStyle

        story = []
        story.append(Paragraph("Relaciones entre Personajes", styles["CustomHeading1"]))

        if not data.relationships:
            story.append(Paragraph("No se detectaron relaciones.", styles["Normal"]))
            return story

        table_data = [["Personaje 1", "Relacion", "Personaje 2"]]
        for rel in data.relationships[:20]:  # Limitar
            table_data.append(
                [
                    rel.get("source_entity_name", ""),
                    rel.get("relation_type", "").replace("_", " ").title(),
                    rel.get("target_entity_name", ""),
                ]
            )

        table = Table(table_data, colWidths=[150, 100, 150])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 1, black),
                ]
            )
        )

        story.append(table)
        return story

    def _build_pdf_style_guide(self, data: ProjectExportData, styles) -> list:
        """Construye seccion de guia de estilo para PDF."""
        from reportlab.platypus import Paragraph

        story = []
        story.append(Paragraph("Guia de Estilo", styles["CustomHeading1"]))

        if not data.style_guide:
            story.append(Paragraph("No se genero guia de estilo.", styles["Normal"]))
            return story

        sg = data.style_guide

        if sg.get("spelling_decisions"):
            story.append(Paragraph("Decisiones de Grafia", styles["CustomHeading2"]))
            for decision in sg["spelling_decisions"]:
                canonical = decision.get("canonical_form", "")
                variants = ", ".join(decision.get("variants", []))
                text = f"<b>{canonical}</b>"
                if variants:
                    text += f" (variantes: {variants})"
                story.append(Paragraph(text, styles["Normal"]))

        if sg.get("style_analysis"):
            sa = sg["style_analysis"]
            story.append(Paragraph("Analisis Estilistico", styles["CustomHeading2"]))

            if sa.get("dialogue_style"):
                story.append(
                    Paragraph(
                        f"<b>Estilo de dialogos:</b> {sa['dialogue_style']}", styles["Normal"]
                    )
                )

            if sa.get("recommendations"):
                story.append(Paragraph("Recomendaciones:", styles["Normal"]))
                for rec in sa["recommendations"]:
                    story.append(Paragraph(f"- {rec}", styles["Normal"]))

        return story


def collect_export_data(
    project_id: int,
    project_manager,
    entity_repository,
    alert_repository,
    chapter_repository,
    options: ExportOptions | None = None,
) -> Result[ProjectExportData]:
    """
    Recopila todos los datos necesarios para exportar un proyecto.

    Args:
        project_id: ID del proyecto
        project_manager: Manager de proyectos
        entity_repository: Repositorio de entidades
        alert_repository: Repositorio de alertas
        chapter_repository: Repositorio de capitulos
        options: Opciones de exportacion

    Returns:
        Result con ProjectExportData
    """
    if options is None:
        options = ExportOptions()

    try:
        # Obtener proyecto
        result = project_manager.get(project_id)
        if result.is_failure:
            return Result.failure(result.error)

        project = result.value

        data = ProjectExportData(
            project_id=project_id,
            project_name=project.name,
            description=project.description or "",
            word_count=project.word_count,
            chapter_count=project.chapter_count,
        )

        # Entidades/personajes
        if options.include_character_sheets and entity_repository:
            from ..entities.models import EntityType

            entities = entity_repository.get_by_project(project_id)
            data.entity_count = len(entities)

            # Solo personajes
            characters = [e for e in entities if e.entity_type == EntityType.CHARACTER]

            # Convertir a diccionarios
            for char in characters:
                char_dict = {
                    "entity_id": char.id,
                    "canonical_name": char.canonical_name,
                    "aliases": char.aliases or [],
                    "entity_type": char.entity_type.value,
                    "importance": char.importance.value
                    if hasattr(char.importance, "value")
                    else str(char.importance),
                    "mention_count": char.mention_count,
                }

                # Obtener atributos
                attrs = entity_repository.get_attributes_by_entity(char.id)
                physical = []
                psychological = []

                for attr in attrs:
                    attr_info = {
                        "key": attr.get("attribute_key", attr.get("name", "")),
                        "value": attr.get("attribute_value", attr.get("value", "")),
                        "confidence": attr.get("confidence", 0.5),
                    }
                    category = attr.get("attribute_type", attr.get("category", ""))
                    if category == "physical":
                        physical.append(attr_info)
                    elif category == "psychological":
                        psychological.append(attr_info)

                char_dict["physical_attributes"] = physical
                char_dict["psychological_attributes"] = psychological

                # Menciones
                char_dict["mentions"] = {
                    "total_mentions": char.mention_count,
                    "chapters": [],
                }

                data.characters.append(char_dict)

        # Alertas
        if options.include_alerts and alert_repository:
            alerts = alert_repository.get_by_project(project_id)
            data.alert_count = len(alerts)

            for alert in alerts:
                alert_dict = {
                    "id": alert.id,
                    "category": alert.category.value
                    if hasattr(alert.category, "value")
                    else str(alert.category),
                    "severity": alert.severity.value
                    if hasattr(alert.severity, "value")
                    else str(alert.severity),
                    "alert_type": alert.alert_type,
                    "title": alert.title,
                    "description": alert.description,
                    "explanation": alert.explanation,
                    "suggestion": alert.suggestion,
                    "chapter": alert.chapter,
                    "scene": getattr(alert, "scene", None),
                    "excerpt": alert.excerpt,
                    "confidence": alert.confidence,
                    "status": alert.status.value
                    if hasattr(alert.status, "value")
                    else str(alert.status),
                }
                data.alerts.append(alert_dict)

        # Timeline
        if options.include_timeline and chapter_repository:
            try:
                from ..temporal import TemporalMarkerExtractor, TimelineBuilder

                chapters = chapter_repository.get_by_project(project_id)

                if chapters:
                    marker_extractor = TemporalMarkerExtractor()
                    all_markers = []

                    for chapter in chapters:
                        chapter_markers = marker_extractor.extract(
                            text=chapter.content,
                            chapter=chapter.chapter_number,
                        )
                        all_markers.extend(chapter_markers)

                    builder = TimelineBuilder()
                    chapter_data = [
                        {
                            "number": ch.chapter_number,
                            "title": ch.title or f"Capitulo {ch.chapter_number}",
                            "start_position": ch.start_char,
                            "content": ch.content,
                        }
                        for ch in chapters
                    ]

                    builder.build_from_markers(all_markers, chapter_data)
                    json_data = builder.export_to_json()
                    data.timeline_events = json_data.get("events", [])

            except ImportError:
                logger.debug("Temporal module not available for export")

        # Relaciones
        if options.include_relationships:
            try:
                from ..relationships.repository import RelationshipRepository

                rel_repo = RelationshipRepository()
                relationships = rel_repo.get_by_project(project_id)

                for rel in relationships:
                    data.relationships.append(rel.to_dict())

            except (ImportError, Exception) as e:
                logger.debug(f"Relationships not available for export: {e}")

        # Style guide
        if options.include_style_guide:
            try:
                from .style_guide import generate_style_guide

                # Obtener texto para analisis
                full_text = ""
                if chapter_repository:
                    chapters = chapter_repository.get_by_project(project_id)
                    full_text = "\n\n".join(ch.content for ch in chapters if ch.content)

                sg_result = generate_style_guide(
                    project_id=project_id,
                    project_name=project.name,
                    text=full_text,
                )

                if sg_result.is_success:
                    data.style_guide = sg_result.value.to_dict()

            except Exception as e:
                logger.debug(f"Style guide not available for export: {e}")

        return Result.success(data)

    except Exception as e:
        error = NarrativeError(
            message=f"Failed to collect export data: {str(e)}",
            severity=ErrorSeverity.ERROR,
            user_message=f"Error recopilando datos para exportar: {str(e)}",
        )
        logger.error(f"Error collecting export data: {e}", exc_info=True)
        return Result.failure(error)
