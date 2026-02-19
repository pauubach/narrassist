"""
Exportador de documento con correcciones como Track Changes.

Genera un documento Word (.docx) que preserva el formato original
pero incluye las correcciones como revisiones de Word (Track Changes),
permitiendo al autor aceptar/rechazar cada cambio individualmente.
"""

import io
import logging
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn

from ..core.errors import ErrorSeverity, NarrativeError
from ..core.result import Result
from ..corrections.base import CorrectionIssue

# Type alias: docx.Document() returns docx.document.Document but the function
# itself is typed as a callable, so we use Any for parameter annotations.
_DocxDocument = Any

logger = logging.getLogger(__name__)


# Namespace para OpenXML de Word
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_PREFIX = "{" + WORD_NS + "}"


@dataclass
class TrackChangeOptions:
    """Opciones para exportación con Track Changes."""

    # Autor de las revisiones
    author: str = "Narrative Assistant"

    # Incluir comentarios con explicación
    include_comments: bool = True

    # Solo aplicar correcciones con confianza alta
    min_confidence: float = 0.5

    # Categorías a incluir (None = todas)
    categories: list[str] | None = None

    # Marcar como revisión (True) o aplicar directamente (False)
    as_track_changes: bool = True


class CorrectedDocumentExporter:
    """
    Exportador que aplica correcciones como Track Changes de Word.

    Preserva el formato original del documento y añade las correcciones
    como revisiones que el autor puede aceptar o rechazar.

    Uso:
        exporter = CorrectedDocumentExporter()
        result = exporter.export(
            original_path="documento.docx",
            corrections=lista_de_correcciones,
            options=TrackChangeOptions()
        )
        with open("documento_corregido.docx", "wb") as f:
            f.write(result.value)
    """

    def __init__(self):
        """Inicializa el exportador."""
        self._revision_id = 0

    def export(
        self,
        original_path: Path,
        corrections: list[CorrectionIssue],
        options: TrackChangeOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta el documento con correcciones como Track Changes.

        Args:
            original_path: Ruta al documento original (.docx)
            corrections: Lista de correcciones a aplicar
            options: Opciones de exportación

        Returns:
            Result con bytes del documento corregido
        """
        if options is None:
            options = TrackChangeOptions()

        try:
            # Validar que es un DOCX
            if not original_path.exists():
                return Result.failure(
                    NarrativeError(
                        f"El archivo no existe: {original_path}",
                        severity=ErrorSeverity.RECOVERABLE,
                    )
                )

            if original_path.suffix.lower() != ".docx":
                return Result.failure(
                    NarrativeError(
                        "Solo se admiten archivos .docx para Track Changes",
                        severity=ErrorSeverity.RECOVERABLE,
                    )
                )

            # Cargar documento
            doc = Document(str(original_path))

            # Filtrar correcciones
            filtered = self._filter_corrections(corrections, options)

            if not filtered:
                logger.info("No hay correcciones que aplicar")
                # Devolver el documento sin cambios
                buffer = io.BytesIO()
                doc.save(buffer)
                return Result.success(buffer.getvalue())

            logger.info(f"Aplicando {len(filtered)} correcciones como Track Changes")

            # Habilitar Track Changes en el documento
            self._enable_track_changes(doc)

            # Aplicar correcciones
            applied_count = self._apply_corrections(doc, filtered, options)

            logger.info(f"Se aplicaron {applied_count} correcciones")

            # Guardar a bytes
            buffer = io.BytesIO()
            doc.save(buffer)

            return Result.success(buffer.getvalue())

        except Exception as e:
            logger.exception(f"Error al exportar con Track Changes: {e}")
            return Result.failure(
                NarrativeError(
                    f"Error al exportar: {str(e)}",
                    severity=ErrorSeverity.RECOVERABLE,
                )
            )

    def export_from_bytes(
        self,
        original_bytes: bytes,
        corrections: list[CorrectionIssue],
        options: TrackChangeOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta desde bytes del documento original.

        Args:
            original_bytes: Contenido del documento original
            corrections: Lista de correcciones a aplicar
            options: Opciones de exportación

        Returns:
            Result con bytes del documento corregido
        """
        if options is None:
            options = TrackChangeOptions()

        try:
            # Cargar documento desde bytes
            doc = Document(io.BytesIO(original_bytes))

            # Filtrar correcciones
            filtered = self._filter_corrections(corrections, options)

            if not filtered:
                logger.info("No hay correcciones que aplicar")
                return Result.success(original_bytes)

            logger.info(f"Aplicando {len(filtered)} correcciones como Track Changes")

            # Habilitar Track Changes
            self._enable_track_changes(doc)

            # Aplicar correcciones
            applied_count = self._apply_corrections(doc, filtered, options)

            logger.info(f"Se aplicaron {applied_count} correcciones")

            # Guardar a bytes
            buffer = io.BytesIO()
            doc.save(buffer)

            return Result.success(buffer.getvalue())

        except Exception as e:
            logger.exception(f"Error al exportar con Track Changes: {e}")
            return Result.failure(
                NarrativeError(
                    f"Error al exportar: {str(e)}",
                    severity=ErrorSeverity.RECOVERABLE,
                )
            )

    def _filter_corrections(
        self,
        corrections: list[CorrectionIssue],
        options: TrackChangeOptions,
    ) -> list[CorrectionIssue]:
        """Filtra correcciones según las opciones."""
        filtered = []

        for c in corrections:
            # Filtrar por confianza
            if c.confidence < options.min_confidence:
                continue

            # Filtrar por categoría
            if options.categories and c.category not in options.categories:
                continue

            # Solo incluir correcciones con sugerencia
            if not c.suggestion:
                continue

            filtered.append(c)

        # Ordenar por posición (de mayor a menor para no afectar índices)
        filtered.sort(key=lambda x: x.start_char, reverse=True)

        return filtered

    def _enable_track_changes(self, doc: _DocxDocument) -> None:
        """Habilita Track Changes en el documento."""
        # Obtener o crear element settings
        settings = doc.settings.element

        # Añadir trackRevisions si no existe
        track_revisions = settings.find(qn("w:trackRevisions"))
        if track_revisions is None:
            track_revisions = parse_xml(f'<w:trackRevisions {nsmap["w"]} w:val="true"/>')
            settings.append(track_revisions)

    def _apply_corrections(
        self,
        doc: _DocxDocument,
        corrections: list[CorrectionIssue],
        options: TrackChangeOptions,
    ) -> int:
        """
        Aplica las correcciones al documento.

        Estrategia: Buscar el texto en los párrafos y aplicar la corrección.
        """
        applied = 0
        self._revision_id = 0

        # Construir mapa de texto del documento
        text_map = self._build_text_map(doc)

        for correction in corrections:
            try:
                if self._apply_single_correction(doc, correction, text_map, options):
                    applied += 1
            except Exception as e:
                logger.warning(
                    f"No se pudo aplicar corrección en posición {correction.start_char}: {e}"
                )

        return applied

    def _build_text_map(self, doc: _DocxDocument) -> list[dict]:
        """
        Construye un mapa de texto con posiciones.

        Returns:
            Lista de {paragraph, run, start, end, text}
        """
        text_map = []
        current_pos = 0

        for para_idx, para in enumerate(doc.paragraphs):
            for run_idx, run in enumerate(para.runs):
                run_text = run.text
                if run_text:
                    text_map.append(
                        {
                            "paragraph_idx": para_idx,
                            "paragraph": para,
                            "run_idx": run_idx,
                            "run": run,
                            "start": current_pos,
                            "end": current_pos + len(run_text),
                            "text": run_text,
                        }
                    )
                    current_pos += len(run_text)

            # Añadir salto de línea entre párrafos
            current_pos += 1  # \n

        return text_map

    def _apply_single_correction(
        self,
        doc: _DocxDocument,
        correction: CorrectionIssue,
        text_map: list[dict],
        options: TrackChangeOptions,
    ) -> bool:
        """
        Aplica una corrección individual.

        Returns:
            True si se aplicó correctamente
        """
        # Buscar el run que contiene el texto
        target_run = None
        target_entry = None

        for entry in text_map:
            if entry["start"] <= correction.start_char < entry["end"]:
                target_run = entry["run"]
                target_entry = entry
                break

        if target_run is None or target_entry is None:
            # Intentar búsqueda por texto literal
            return self._apply_by_text_search(doc, correction, options)

        # Calcular posición relativa dentro del run
        relative_start = correction.start_char - target_entry["start"]
        relative_end = min(correction.end_char - target_entry["start"], len(target_entry["text"]))

        original_text = target_entry["text"]

        # Si la corrección abarca todo el run, es simple
        if relative_start == 0 and relative_end >= len(original_text):
            if options.as_track_changes:
                self._apply_track_change_to_run(
                    target_run,
                    original_text,
                    correction.suggestion or "",
                    options.author,
                )
            else:
                target_run.text = correction.suggestion or ""
            return True

        # Corrección parcial: dividir el run
        return self._apply_partial_correction(
            target_run,
            original_text,
            relative_start,
            relative_end,
            correction.suggestion or "",
            options,
        )

    def _apply_by_text_search(
        self,
        doc: _DocxDocument,
        correction: CorrectionIssue,
        options: TrackChangeOptions,
    ) -> bool:
        """
        Aplica corrección buscando el texto literal.

        Fallback cuando las posiciones no coinciden exactamente.
        """
        search_text = correction.text

        for para in doc.paragraphs:
            if search_text in para.text:
                # Encontrar el run específico
                for run in para.runs:
                    if search_text in run.text:
                        if options.as_track_changes:
                            # Reemplazar con track change
                            run.text = run.text.replace(
                                search_text,
                                correction.suggestion or "",
                                1,  # Solo primera ocurrencia
                            )
                            self._apply_track_change_to_run(
                                run,
                                search_text,
                                correction.suggestion or "",
                                options.author,
                            )
                        else:
                            run.text = run.text.replace(search_text, correction.suggestion or "", 1)
                        return True

        return False

    def _apply_track_change_to_run(
        self,
        run,
        old_text: str,
        new_text: str,
        author: str,
    ) -> None:
        """
        Aplica Track Change a un run completo.

        Crea elementos <w:del> para texto eliminado y <w:ins> para insertado.
        """
        self._revision_id += 1
        revision_id = self._revision_id
        date = datetime.now().isoformat()

        # Obtener el elemento padre del run
        run_element = run._element
        parent = run_element.getparent()

        # Crear el texto de eliminación (del)
        if old_text:
            del_element = parse_xml(
                f'<w:del xmlns:w="{WORD_NS}" w:id="{revision_id}" '
                f'w:author="{author}" w:date="{date}">'
                f"<w:r><w:delText>{self._escape_xml(old_text)}</w:delText></w:r>"
                f"</w:del>"
            )
            parent.insert(list(parent).index(run_element), del_element)

        # Crear el texto de inserción (ins)
        if new_text:
            self._revision_id += 1
            ins_element = parse_xml(
                f'<w:ins xmlns:w="{WORD_NS}" w:id="{self._revision_id}" '
                f'w:author="{author}" w:date="{date}">'
                f"<w:r><w:t>{self._escape_xml(new_text)}</w:t></w:r>"
                f"</w:ins>"
            )
            parent.insert(list(parent).index(run_element), ins_element)

        # Eliminar el run original
        parent.remove(run_element)

    def _apply_partial_correction(
        self,
        run,
        original_text: str,
        start: int,
        end: int,
        replacement: str,
        options: TrackChangeOptions,
    ) -> bool:
        """
        Aplica una corrección parcial dividiendo el run.

        El texto se divide en: [antes][corregido][después]
        """
        before = original_text[:start]
        target = original_text[start:end]
        after = original_text[end:]

        run_element = run._element
        parent = run_element.getparent()
        run_index = list(parent).index(run_element)

        # Preservar formato del run original
        run_props = run_element.find(qn("w:rPr"))

        # Crear run para texto antes (si existe)
        if before:
            before_run = parse_xml(
                f'<w:r xmlns:w="{WORD_NS}">'
                f'<w:t xml:space="preserve">{self._escape_xml(before)}</w:t>'
                f"</w:r>"
            )
            if run_props is not None:
                before_run.insert(0, deepcopy(run_props))
            parent.insert(run_index, before_run)
            run_index += 1

        # Aplicar track change para el texto objetivo
        if options.as_track_changes:
            self._revision_id += 1
            date = datetime.now().isoformat()

            # Eliminación
            del_element = parse_xml(
                f'<w:del xmlns:w="{WORD_NS}" w:id="{self._revision_id}" '
                f'w:author="{options.author}" w:date="{date}">'
                f"<w:r><w:delText>{self._escape_xml(target)}</w:delText></w:r>"
                f"</w:del>"
            )
            parent.insert(run_index, del_element)
            run_index += 1

            # Inserción
            if replacement:
                self._revision_id += 1
                ins_element = parse_xml(
                    f'<w:ins xmlns:w="{WORD_NS}" w:id="{self._revision_id}" '
                    f'w:author="{options.author}" w:date="{date}">'
                    f"<w:r><w:t>{self._escape_xml(replacement)}</w:t></w:r>"
                    f"</w:ins>"
                )
                parent.insert(run_index, ins_element)
                run_index += 1
        else:
            # Aplicar directamente sin track changes
            if replacement:
                replace_run = parse_xml(
                    f'<w:r xmlns:w="{WORD_NS}"><w:t>{self._escape_xml(replacement)}</w:t></w:r>'
                )
                if run_props is not None:
                    replace_run.insert(0, deepcopy(run_props))
                parent.insert(run_index, replace_run)
                run_index += 1

        # Crear run para texto después (si existe)
        if after:
            after_run = parse_xml(
                f'<w:r xmlns:w="{WORD_NS}">'
                f'<w:t xml:space="preserve">{self._escape_xml(after)}</w:t>'
                f"</w:r>"
            )
            if run_props is not None:
                after_run.insert(0, deepcopy(run_props))
            parent.insert(run_index, after_run)

        # Eliminar run original
        parent.remove(run_element)

        return True

    def _escape_xml(self, text: str) -> str:
        """Escapa caracteres especiales para XML."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )


def export_with_track_changes(
    original_path: Path,
    corrections: list[CorrectionIssue],
    output_path: Path | None = None,
    options: TrackChangeOptions | None = None,
) -> Result[Path]:
    """
    Función de conveniencia para exportar con Track Changes.

    Args:
        original_path: Ruta al documento original
        corrections: Lista de correcciones
        output_path: Ruta de salida (default: original_corregido.docx)
        options: Opciones de exportación

    Returns:
        Result con la ruta del archivo generado
    """
    exporter = CorrectedDocumentExporter()

    result = exporter.export(original_path, corrections, options)

    if result.is_failure:
        return Result.failure(result.error)

    # Determinar ruta de salida
    if output_path is None:
        output_path = original_path.with_name(
            f"{original_path.stem}_corregido{original_path.suffix}"
        )

    # Guardar archivo
    try:
        with open(output_path, "wb") as f:
            assert result.value is not None
            f.write(result.value)
        logger.info(f"Documento corregido guardado en: {output_path}")
        return Result.success(output_path)
    except Exception as e:
        return Result.failure(
            NarrativeError(
                f"Error al guardar archivo: {e}",
                severity=ErrorSeverity.RECOVERABLE,
            )
        )
