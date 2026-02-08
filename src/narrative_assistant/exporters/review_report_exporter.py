"""
Exportador de informes de revisión editorial.

Genera informes detallados en PDF y DOCX con estadísticas
de los problemas detectados por los 14 detectores de corrección.

Incluye:
- Resumen ejecutivo con totales por categoría
- Distribución de problemas por confianza
- Listado detallado por categoría
- Gráficos de distribución (PDF)
- Recomendaciones de estilo
"""

import io
import logging
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..core.errors import ErrorSeverity, NarrativeError
from ..core.result import Result
from ..corrections.base import CorrectionIssue

logger = logging.getLogger(__name__)


# Nombres amigables para categorías
CATEGORY_DISPLAY_NAMES = {
    "typography": "Tipografía",
    "repetition": "Repeticiones",
    "agreement": "Concordancia",
    "punctuation": "Puntuación",
    "terminology": "Terminología",
    "regional": "Regionalismos",
    "clarity": "Claridad",
    "grammar": "Gramática",
    "anglicisms": "Extranjerismos",
    "crutch_words": "Muletillas",
    "glossary": "Glosario",
    "anacoluto": "Anacolutos",
    "pov": "Punto de vista",
    "orthography": "Ortografía RAE",
}

# Colores por categoría (para visualización)
CATEGORY_COLORS = {
    "typography": RGBColor(41, 128, 185),  # Azul
    "repetition": RGBColor(155, 89, 182),  # Morado
    "agreement": RGBColor(231, 76, 60),  # Rojo
    "punctuation": RGBColor(230, 126, 34),  # Naranja
    "terminology": RGBColor(46, 204, 113),  # Verde
    "regional": RGBColor(26, 188, 156),  # Turquesa
    "clarity": RGBColor(52, 73, 94),  # Gris oscuro
    "grammar": RGBColor(192, 57, 43),  # Rojo oscuro
    "anglicisms": RGBColor(241, 196, 15),  # Amarillo
    "crutch_words": RGBColor(142, 68, 173),  # Púrpura
    "glossary": RGBColor(39, 174, 96),  # Verde esmeralda
    "anacoluto": RGBColor(211, 84, 0),  # Naranja oscuro
    "pov": RGBColor(127, 140, 141),  # Gris
    "orthography": RGBColor(44, 62, 80),  # Azul oscuro
}


@dataclass
class ReviewReportOptions:
    """Opciones para el informe de revisión."""

    # Título del documento
    document_title: str = "Documento"

    # Incluir secciones
    include_summary: bool = True
    include_by_category: bool = True
    include_by_chapter: bool = True
    include_detailed_list: bool = True
    include_recommendations: bool = True

    # Filtros
    min_confidence: float = 0.0
    categories: list[str] | None = None
    max_issues_per_category: int = 50

    # Formato
    include_context: bool = True
    include_suggestions: bool = True
    sort_by: str = "confidence"  # confidence, position, category


@dataclass
class CategoryStats:
    """Estadísticas por categoría."""

    category: str
    display_name: str
    total: int = 0
    high_confidence: int = 0  # >= 0.8
    medium_confidence: int = 0  # 0.5 - 0.8
    low_confidence: int = 0  # < 0.5
    by_type: dict[str, int] = field(default_factory=dict)
    issues: list[CorrectionIssue] = field(default_factory=list)


@dataclass
class ReviewReportData:
    """Datos completos del informe de revisión."""

    document_title: str
    generated_at: datetime
    total_issues: int
    total_by_confidence: dict[str, int]
    categories: list[CategoryStats]
    by_chapter: dict[int, int]
    top_issues_by_type: list[tuple[str, str, int]]  # (category, type, count)


class ReviewReportExporter:
    """
    Exportador de informes de revisión editorial.

    Genera informes detallados de los problemas detectados por
    los 14 detectores de corrección editorial.

    Uso:
        exporter = ReviewReportExporter()
        result = exporter.export_to_docx(issues, options)
        # o
        result = exporter.export_to_pdf(issues, options)
    """

    def __init__(self):
        """Inicializa el exportador."""
        pass

    def prepare_report_data(
        self,
        issues: list[CorrectionIssue],
        options: ReviewReportOptions | None = None,
    ) -> ReviewReportData:
        """
        Prepara los datos para el informe.

        Args:
            issues: Lista de CorrectionIssue
            options: Opciones del informe

        Returns:
            ReviewReportData con estadísticas calculadas
        """
        if options is None:
            options = ReviewReportOptions()

        # Filtrar issues
        filtered = self._filter_issues(issues, options)

        # Agrupar por categoría
        by_category: dict[str, list[CorrectionIssue]] = defaultdict(list)
        for issue in filtered:
            by_category[issue.category].append(issue)

        # Calcular estadísticas por categoría
        categories_stats: list[CategoryStats] = []
        for cat_name in sorted(by_category.keys()):
            cat_issues = by_category[cat_name]
            stats = CategoryStats(
                category=cat_name,
                display_name=CATEGORY_DISPLAY_NAMES.get(cat_name, cat_name.title()),
                total=len(cat_issues),
                issues=cat_issues[: options.max_issues_per_category],
            )

            # Contar por confianza
            for issue in cat_issues:
                if issue.confidence >= 0.8:
                    stats.high_confidence += 1
                elif issue.confidence >= 0.5:
                    stats.medium_confidence += 1
                else:
                    stats.low_confidence += 1

                # Contar por tipo
                stats.by_type[issue.issue_type] = stats.by_type.get(issue.issue_type, 0) + 1

            categories_stats.append(stats)

        # Ordenar categorías por total descendente
        categories_stats.sort(key=lambda x: x.total, reverse=True)

        # Agrupar por capítulo
        by_chapter: dict[int, int] = Counter()
        for issue in filtered:
            chapter = issue.chapter_index or 0
            by_chapter[chapter] += 1

        # Calcular totales por confianza
        total_by_confidence = {
            "high": sum(1 for i in filtered if i.confidence >= 0.8),
            "medium": sum(1 for i in filtered if 0.5 <= i.confidence < 0.8),
            "low": sum(1 for i in filtered if i.confidence < 0.5),
        }

        # Top issues por tipo
        type_counts: Counter = Counter()
        for issue in filtered:
            type_counts[(issue.category, issue.issue_type)] += 1
        top_issues = [(cat, typ, count) for (cat, typ), count in type_counts.most_common(10)]

        return ReviewReportData(
            document_title=options.document_title,
            generated_at=datetime.now(),
            total_issues=len(filtered),
            total_by_confidence=total_by_confidence,
            categories=categories_stats,
            by_chapter=dict(by_chapter),
            top_issues_by_type=top_issues,
        )

    def _filter_issues(
        self,
        issues: list[CorrectionIssue],
        options: ReviewReportOptions,
    ) -> list[CorrectionIssue]:
        """Filtra issues según opciones."""
        filtered = []

        for issue in issues:
            # Filtrar por confianza
            if issue.confidence < options.min_confidence:
                continue

            # Filtrar por categoría
            if options.categories and issue.category not in options.categories:
                continue

            filtered.append(issue)

        # Ordenar
        if options.sort_by == "confidence":
            filtered.sort(key=lambda x: x.confidence, reverse=True)
        elif options.sort_by == "position":
            filtered.sort(key=lambda x: x.start_char)
        elif options.sort_by == "category":
            filtered.sort(key=lambda x: (x.category, -x.confidence))

        return filtered

    def export_to_docx(
        self,
        issues: list[CorrectionIssue],
        options: ReviewReportOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta el informe a formato DOCX.

        Args:
            issues: Lista de CorrectionIssue
            options: Opciones del informe

        Returns:
            Result con bytes del documento DOCX
        """
        if options is None:
            options = ReviewReportOptions()

        try:
            # Preparar datos
            data = self.prepare_report_data(issues, options)

            # Crear documento
            doc = Document()
            self._setup_docx_styles(doc)

            # Portada
            self._add_docx_cover(doc, data)

            # Resumen ejecutivo
            if options.include_summary:
                self._add_docx_summary(doc, data)

            # Desglose por categoría
            if options.include_by_category:
                self._add_docx_by_category(doc, data)

            # Desglose por capítulo
            if options.include_by_chapter and data.by_chapter:
                self._add_docx_by_chapter(doc, data)

            # Listado detallado
            if options.include_detailed_list:
                self._add_docx_detailed_list(doc, data, options)

            # Recomendaciones
            if options.include_recommendations:
                self._add_docx_recommendations(doc, data)

            # Footer
            self._add_docx_footer(doc, data)

            # Guardar a bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(f"Exported review report DOCX: {data.total_issues} issues")
            return Result.success(buffer.getvalue())

        except Exception as e:
            error = NarrativeError(
                message=f"Failed to export review report DOCX: {str(e)}",
                severity=ErrorSeverity.ERROR,
                user_message=f"Error exportando informe DOCX: {str(e)}",
            )
            logger.error(f"Error exporting review report DOCX: {e}", exc_info=True)
            return Result.failure(error)

    def export_to_pdf(
        self,
        issues: list[CorrectionIssue],
        options: ReviewReportOptions | None = None,
    ) -> Result[bytes]:
        """
        Exporta el informe a formato PDF.

        Args:
            issues: Lista de CorrectionIssue
            options: Opciones del informe

        Returns:
            Result con bytes del documento PDF
        """
        if options is None:
            options = ReviewReportOptions()

        try:
            # Intentar importar reportlab
            from reportlab.lib.colors import HexColor, black, gray, lightgrey
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                ListFlowable,
                ListItem,
                PageBreak,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )

            # Preparar datos
            data = self.prepare_report_data(issues, options)

            buffer = io.BytesIO()

            # Crear documento PDF
            pdf_doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=2 * cm,
                leftMargin=2 * cm,
                topMargin=2 * cm,
                bottomMargin=2 * cm,
            )

            # Configurar estilos
            styles = getSampleStyleSheet()
            self._setup_pdf_styles(styles)

            story = []

            # Portada
            story.extend(self._build_pdf_cover(data, styles))
            story.append(PageBreak())

            # Resumen ejecutivo
            if options.include_summary:
                story.extend(self._build_pdf_summary(data, styles))

            # Desglose por categoría
            if options.include_by_category:
                story.extend(self._build_pdf_by_category(data, styles))

            # Desglose por capítulo
            if options.include_by_chapter and data.by_chapter:
                story.extend(self._build_pdf_by_chapter(data, styles))

            # Listado detallado (resumido para PDF)
            if options.include_detailed_list:
                story.extend(self._build_pdf_detailed_list(data, options, styles))

            # Recomendaciones
            if options.include_recommendations:
                story.extend(self._build_pdf_recommendations(data, styles))

            # Footer
            story.append(Spacer(1, 30))
            story.append(
                Paragraph(
                    f"Generado por Narrative Assistant el {data.generated_at.strftime('%Y-%m-%d %H:%M')}",
                    styles["Normal"],
                )
            )

            # Construir PDF
            pdf_doc.build(story)
            buffer.seek(0)

            logger.info(f"Exported review report PDF: {data.total_issues} issues")
            return Result.success(buffer.getvalue())

        except ImportError:
            error = NarrativeError(
                message="reportlab not installed",
                severity=ErrorSeverity.ERROR,
                user_message="Para exportar a PDF es necesario instalar reportlab: pip install reportlab",
            )
            return Result.failure(error)
        except Exception as e:
            error = NarrativeError(
                message=f"Failed to export review report PDF: {str(e)}",
                severity=ErrorSeverity.ERROR,
                user_message=f"Error exportando informe PDF: {str(e)}",
            )
            logger.error(f"Error exporting review report PDF: {e}", exc_info=True)
            return Result.failure(error)

    # =========================================================================
    # DOCX Generation Methods
    # =========================================================================

    def _setup_docx_styles(self, doc: Document) -> None:
        """Configura estilos del documento DOCX."""
        styles = doc.styles

        # Estilo para título de sección
        try:
            section_style = styles.add_style("ReviewSection", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(26, 26, 46)
            section_style.paragraph_format.space_before = Pt(20)
            section_style.paragraph_format.space_after = Pt(10)
        except ValueError:
            pass

        # Estilo para subsección
        try:
            subsection_style = styles.add_style("ReviewSubsection", WD_STYLE_TYPE.PARAGRAPH)
            subsection_style.font.size = Pt(13)
            subsection_style.font.bold = True
            subsection_style.font.color.rgb = RGBColor(45, 45, 74)
            subsection_style.paragraph_format.space_before = Pt(15)
            subsection_style.paragraph_format.space_after = Pt(8)
        except ValueError:
            pass

        # Estilo para contexto/cita
        try:
            context_style = styles.add_style("ReviewContext", WD_STYLE_TYPE.PARAGRAPH)
            context_style.font.size = Pt(10)
            context_style.font.italic = True
            context_style.font.color.rgb = RGBColor(100, 100, 100)
            context_style.paragraph_format.left_indent = Inches(0.3)
        except ValueError:
            pass

        # Estilo para número de estadística
        try:
            stat_style = styles.add_style("StatNumber", WD_STYLE_TYPE.PARAGRAPH)
            stat_style.font.size = Pt(28)
            stat_style.font.bold = True
            stat_style.font.color.rgb = RGBColor(41, 128, 185)
        except ValueError:
            pass

    def _add_docx_cover(self, doc: Document, data: ReviewReportData) -> None:
        """Añade portada al documento DOCX."""
        # Espaciado superior
        for _ in range(3):
            doc.add_paragraph()

        # Título
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Informe de Revisión Editorial")
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(26, 26, 46)

        # Subtítulo
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(data.document_title)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(74, 74, 106)

        # Espaciado
        for _ in range(2):
            doc.add_paragraph()

        # Estadísticas destacadas
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = stats_para.add_run(f"{data.total_issues}")
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = RGBColor(41, 128, 185)

        label_para = doc.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label_para.add_run("observaciones detectadas")
        run.font.size = Pt(14)

        # Espaciado
        for _ in range(3):
            doc.add_paragraph()

        # Fecha
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(
            f"Generado el {data.generated_at.strftime('%d de %B de %Y a las %H:%M')}"
        )
        run.font.size = Pt(11)
        run.font.italic = True

        doc.add_page_break()

    def _add_docx_summary(self, doc: Document, data: ReviewReportData) -> None:
        """Añade resumen ejecutivo al DOCX."""
        doc.add_heading("Resumen Ejecutivo", level=1)

        # Tabla de totales
        p = doc.add_paragraph()
        p.add_run("Total de observaciones: ").bold = True
        p.add_run(f"{data.total_issues}")

        # Distribución por confianza
        doc.add_heading("Distribución por Confianza", level=2)

        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"

        headers = ["Nivel", "Cantidad", "Porcentaje"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        confidence_data = [
            ("Alta (≥80%)", data.total_by_confidence["high"]),
            ("Media (50-80%)", data.total_by_confidence["medium"]),
            ("Baja (<50%)", data.total_by_confidence["low"]),
        ]

        for i, (level, count) in enumerate(confidence_data, 1):
            table.rows[i].cells[0].text = level
            table.rows[i].cells[1].text = str(count)
            pct = (count / data.total_issues * 100) if data.total_issues > 0 else 0
            table.rows[i].cells[2].text = f"{pct:.1f}%"

        doc.add_paragraph()

        # Top 10 tipos de issues
        if data.top_issues_by_type:
            doc.add_heading("Tipos de Observaciones más Frecuentes", level=2)

            for i, (category, issue_type, count) in enumerate(data.top_issues_by_type[:10], 1):
                cat_name = CATEGORY_DISPLAY_NAMES.get(category, category)
                type_name = issue_type.replace("_", " ").title()
                p = doc.add_paragraph(style="List Number")
                p.add_run(f"{cat_name} - {type_name}: ").bold = True
                p.add_run(f"{count}")

        doc.add_paragraph()

    def _add_docx_by_category(self, doc: Document, data: ReviewReportData) -> None:
        """Añade desglose por categoría al DOCX."""
        doc.add_heading("Desglose por Categoría", level=1)

        # Tabla de categorías
        table = doc.add_table(rows=len(data.categories) + 1, cols=5)
        table.style = "Table Grid"

        # Encabezados
        headers = ["Categoría", "Total", "Alta", "Media", "Baja"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Datos
        for i, cat in enumerate(data.categories, 1):
            row = table.rows[i]
            row.cells[0].text = cat.display_name
            row.cells[1].text = str(cat.total)
            row.cells[2].text = str(cat.high_confidence)
            row.cells[3].text = str(cat.medium_confidence)
            row.cells[4].text = str(cat.low_confidence)

        doc.add_paragraph()

        # Detalles por categoría
        for cat in data.categories:
            if cat.total == 0:
                continue

            doc.add_heading(f"{cat.display_name} ({cat.total})", level=2)

            # Tipos dentro de la categoría
            if cat.by_type:
                p = doc.add_paragraph()
                p.add_run("Tipos detectados: ").bold = True

                type_items = sorted(cat.by_type.items(), key=lambda x: -x[1])
                type_text = ", ".join(
                    f"{t.replace('_', ' ').title()} ({c})" for t, c in type_items[:5]
                )
                p.add_run(type_text)

            doc.add_paragraph()

    def _add_docx_by_chapter(self, doc: Document, data: ReviewReportData) -> None:
        """Añade desglose por capítulo al DOCX."""
        doc.add_heading("Distribución por Capítulo", level=1)

        # Tabla
        sorted_chapters = sorted(data.by_chapter.items())
        table = doc.add_table(rows=len(sorted_chapters) + 1, cols=2)
        table.style = "Table Grid"

        # Encabezados
        table.rows[0].cells[0].text = "Capítulo"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[1].text = "Observaciones"
        table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

        # Datos
        for i, (chapter, count) in enumerate(sorted_chapters, 1):
            table.rows[i].cells[0].text = f"Capítulo {chapter}" if chapter > 0 else "Sin capítulo"
            table.rows[i].cells[1].text = str(count)

        doc.add_paragraph()

    def _add_docx_detailed_list(
        self,
        doc: Document,
        data: ReviewReportData,
        options: ReviewReportOptions,
    ) -> None:
        """Añade listado detallado al DOCX."""
        doc.add_heading("Listado Detallado de Observaciones", level=1)

        for cat in data.categories:
            if not cat.issues:
                continue

            doc.add_heading(f"{cat.display_name}", level=2)

            for i, issue in enumerate(cat.issues, 1):
                # Título del issue
                p = doc.add_paragraph()
                type_name = issue.issue_type.replace("_", " ").title()
                p.add_run(f"{i}. {type_name}").bold = True
                p.add_run(f" (confianza: {issue.confidence:.0%})")

                # Texto problemático
                p = doc.add_paragraph()
                p.add_run("Texto: ").bold = True
                p.add_run(f'"{issue.text}"')

                # Explicación
                if issue.explanation:
                    p = doc.add_paragraph()
                    p.add_run("Explicación: ").bold = True
                    p.add_run(issue.explanation)

                # Sugerencia
                if options.include_suggestions and issue.suggestion:
                    p = doc.add_paragraph()
                    p.add_run("Sugerencia: ").bold = True
                    p.add_run(f'"{issue.suggestion}"')

                # Contexto
                if options.include_context and issue.context:
                    context_para = doc.add_paragraph()
                    context_para.add_run("Contexto: ").bold = True
                    ctx_run = context_para.add_run(issue.context)
                    ctx_run.italic = True
                    ctx_run.font.size = Pt(10)

                # Ubicación
                if issue.chapter_index:
                    p = doc.add_paragraph()
                    p.add_run(f"Capítulo: {issue.chapter_index}")
                    p.font.size = Pt(9)

                doc.add_paragraph()  # Espaciado

    def _add_docx_recommendations(self, doc: Document, data: ReviewReportData) -> None:
        """Añade recomendaciones al DOCX."""
        doc.add_heading("Recomendaciones", level=1)

        recommendations = self._generate_recommendations(data)

        for rec in recommendations:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(rec)

        doc.add_paragraph()

    def _add_docx_footer(self, doc: Document, data: ReviewReportData) -> None:
        """Añade footer al DOCX."""
        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Informe generado por Narrative Assistant el "
            f"{data.generated_at.strftime('%Y-%m-%d %H:%M')}"
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    # =========================================================================
    # PDF Generation Methods
    # =========================================================================

    def _setup_pdf_styles(self, styles) -> None:
        """Configura estilos adicionales para PDF."""
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import ParagraphStyle

        # Título de portada
        styles.add(
            ParagraphStyle(
                name="ReportTitle",
                parent=styles["Title"],
                fontSize=28,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=HexColor("#1a1a2e"),
            )
        )

        # Subtítulo
        styles.add(
            ParagraphStyle(
                name="ReportSubtitle",
                parent=styles["Normal"],
                fontSize=16,
                spaceAfter=15,
                alignment=TA_CENTER,
                textColor=HexColor("#4a4a6a"),
            )
        )

        # Número grande
        styles.add(
            ParagraphStyle(
                name="BigNumber",
                parent=styles["Normal"],
                fontSize=48,
                alignment=TA_CENTER,
                textColor=HexColor("#2980b9"),
                fontName="Helvetica-Bold",
            )
        )

        # Sección
        styles.add(
            ParagraphStyle(
                name="SectionHeader",
                parent=styles["Heading1"],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=10,
                textColor=HexColor("#1a1a2e"),
            )
        )

        # Subsección
        styles.add(
            ParagraphStyle(
                name="SubsectionHeader",
                parent=styles["Heading2"],
                fontSize=13,
                spaceBefore=15,
                spaceAfter=8,
                textColor=HexColor("#2d2d4a"),
            )
        )

    def _build_pdf_cover(self, data: ReviewReportData, styles) -> list:
        """Construye la portada para PDF."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Spacer(1, 80))
        story.append(Paragraph("Informe de Revisión Editorial", styles["ReportTitle"]))
        story.append(Paragraph(data.document_title, styles["ReportSubtitle"]))
        story.append(Spacer(1, 40))
        story.append(Paragraph(str(data.total_issues), styles["BigNumber"]))
        story.append(Paragraph("observaciones detectadas", styles["ReportSubtitle"]))
        story.append(Spacer(1, 60))
        story.append(
            Paragraph(
                f"Generado el {data.generated_at.strftime('%d de %B de %Y a las %H:%M')}",
                styles["Normal"],
            )
        )

        return story

    def _build_pdf_summary(self, data: ReviewReportData, styles) -> list:
        """Construye resumen ejecutivo para PDF."""
        from reportlab.lib.colors import black, lightgrey
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

        story = []
        story.append(Paragraph("Resumen Ejecutivo", styles["SectionHeader"]))

        # Totales
        story.append(
            Paragraph(f"<b>Total de observaciones:</b> {data.total_issues}", styles["Normal"])
        )
        story.append(Spacer(1, 10))

        # Tabla de confianza
        story.append(Paragraph("Distribución por Confianza", styles["SubsectionHeader"]))

        table_data = [
            ["Nivel", "Cantidad", "Porcentaje"],
            [
                "Alta (≥80%)",
                str(data.total_by_confidence["high"]),
                f"{(data.total_by_confidence['high'] / data.total_issues * 100) if data.total_issues > 0 else 0:.1f}%",
            ],
            [
                "Media (50-80%)",
                str(data.total_by_confidence["medium"]),
                f"{(data.total_by_confidence['medium'] / data.total_issues * 100) if data.total_issues > 0 else 0:.1f}%",
            ],
            [
                "Baja (<50%)",
                str(data.total_by_confidence["low"]),
                f"{(data.total_by_confidence['low'] / data.total_issues * 100) if data.total_issues > 0 else 0:.1f}%",
            ],
        ]

        table = Table(table_data, colWidths=[120, 80, 80])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 1, black),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 15))

        return story

    def _build_pdf_by_category(self, data: ReviewReportData, styles) -> list:
        """Construye desglose por categoría para PDF."""
        from reportlab.lib.colors import black, lightgrey
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

        story = []
        story.append(Paragraph("Desglose por Categoría", styles["SectionHeader"]))

        table_data = [["Categoría", "Total", "Alta", "Media", "Baja"]]
        for cat in data.categories:
            table_data.append(
                [
                    cat.display_name,
                    str(cat.total),
                    str(cat.high_confidence),
                    str(cat.medium_confidence),
                    str(cat.low_confidence),
                ]
            )

        table = Table(table_data, colWidths=[120, 60, 60, 60, 60])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 1, black),
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 20))

        return story

    def _build_pdf_by_chapter(self, data: ReviewReportData, styles) -> list:
        """Construye desglose por capítulo para PDF."""
        from reportlab.lib.colors import black, lightgrey
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

        story = []
        story.append(Paragraph("Distribución por Capítulo", styles["SectionHeader"]))

        table_data = [["Capítulo", "Observaciones"]]
        for chapter, count in sorted(data.by_chapter.items()):
            chapter_name = f"Capítulo {chapter}" if chapter > 0 else "Sin capítulo"
            table_data.append([chapter_name, str(count)])

        table = Table(table_data, colWidths=[150, 100])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), lightgrey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 1, black),
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 20))

        return story

    def _build_pdf_detailed_list(
        self,
        data: ReviewReportData,
        options: ReviewReportOptions,
        styles,
    ) -> list:
        """Construye listado detallado para PDF (resumido)."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Paragraph("Observaciones Destacadas", styles["SectionHeader"]))

        # Solo mostrar las primeras de cada categoría
        for cat in data.categories[:5]:
            if not cat.issues:
                continue

            story.append(
                Paragraph(
                    f"<b>{cat.display_name}</b> ({cat.total} total)", styles["SubsectionHeader"]
                )
            )

            for issue in cat.issues[:3]:
                type_name = issue.issue_type.replace("_", " ").title()
                text = (
                    f'<b>{type_name}</b>: "{issue.text[:50]}..."'
                    if len(issue.text) > 50
                    else f'<b>{type_name}</b>: "{issue.text}"'
                )
                story.append(Paragraph(text, styles["Normal"]))

                if issue.suggestion:
                    story.append(Paragraph(f'→ Sugerencia: "{issue.suggestion}"', styles["Normal"]))

                story.append(Spacer(1, 5))

            story.append(Spacer(1, 10))

        return story

    def _build_pdf_recommendations(self, data: ReviewReportData, styles) -> list:
        """Construye recomendaciones para PDF."""
        from reportlab.platypus import Paragraph, Spacer

        story = []
        story.append(Paragraph("Recomendaciones", styles["SectionHeader"]))

        recommendations = self._generate_recommendations(data)

        for rec in recommendations:
            story.append(Paragraph(f"• {rec}", styles["Normal"]))

        story.append(Spacer(1, 15))

        return story

    # =========================================================================
    # Helper Methods
    # =========================================================================

    def _generate_recommendations(self, data: ReviewReportData) -> list[str]:
        """Genera recomendaciones basadas en los datos del informe."""
        recommendations = []

        # Categoría con más problemas
        if data.categories:
            top_cat = data.categories[0]
            recommendations.append(
                f"Revisar especialmente la categoría '{top_cat.display_name}' "
                f"con {top_cat.total} observaciones detectadas."
            )

        # Nivel de confianza
        high_ratio = (
            data.total_by_confidence["high"] / data.total_issues if data.total_issues > 0 else 0
        )
        if high_ratio > 0.5:
            recommendations.append(
                "La mayoría de observaciones tienen alta confianza (≥80%). "
                "Se recomienda revisarlas prioritariamente."
            )
        elif data.total_by_confidence["low"] > data.total_by_confidence["high"]:
            recommendations.append(
                "Muchas observaciones tienen baja confianza. "
                "Revisar con criterio editorial antes de aplicar cambios."
            )

        # Por tipos específicos
        for category, issue_type, count in data.top_issues_by_type[:3]:
            cat_display = CATEGORY_DISPLAY_NAMES.get(category, category)
            type_name = issue_type.replace("_", " ")

            if issue_type in ("lexical_close", "sentence_start"):
                recommendations.append(
                    f"Se detectaron {count} repeticiones de tipo '{type_name}' ({cat_display}). "
                    "Considerar variar el vocabulario para mayor fluidez."
                )
            elif issue_type in ("wrong_dash_dialogue", "wrong_quote_style"):
                recommendations.append(
                    f"Revisar {count} problemas de '{type_name}' en {cat_display.lower()}."
                )
            elif issue_type == "sentence_too_long":
                recommendations.append(
                    f"Hay {count} oraciones excesivamente largas. "
                    "Considerar dividirlas para mejorar la claridad."
                )

        # Recomendación general
        if data.total_issues > 100:
            recommendations.append(
                "El documento tiene un número elevado de observaciones. "
                "Se sugiere una revisión por fases, comenzando por las de mayor confianza."
            )
        elif data.total_issues < 20:
            recommendations.append(
                "El documento tiene pocas observaciones. "
                "Revisión general recomendada para verificar coherencia."
            )

        return recommendations[:6]  # Máximo 6 recomendaciones


def export_review_report(
    issues: list[CorrectionIssue],
    output_path: Path,
    format: str = "docx",
    options: ReviewReportOptions | None = None,
) -> Result[Path]:
    """
    Función de conveniencia para exportar informe de revisión.

    Args:
        issues: Lista de CorrectionIssue
        output_path: Ruta de salida
        format: Formato de salida ("docx" o "pdf")
        options: Opciones del informe

    Returns:
        Result con la ruta del archivo generado
    """
    exporter = ReviewReportExporter()

    if format.lower() == "pdf":
        result = exporter.export_to_pdf(issues, options)
    else:
        result = exporter.export_to_docx(issues, options)

    if result.is_failure:
        return result

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(result.value)
        logger.info(f"Review report saved to: {output_path}")
        return Result.success(output_path)
    except Exception as e:
        return Result.failure(
            NarrativeError(
                f"Error saving report: {e}",
                severity=ErrorSeverity.ERROR,
            )
        )
