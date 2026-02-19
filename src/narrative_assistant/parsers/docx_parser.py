"""
Parser para documentos Microsoft Word (DOCX).

DOCX es el formato prioritario para el MVP:
- Más común en flujos editoriales
- Preserva estructura (headings, estilos)
- Mejor soporte para detección de capítulos
"""

import logging
import re
from pathlib import Path
from zipfile import BadZipFile

from ..core.errors import CorruptedDocumentError, EmptyDocumentError
from ..core.result import Result
from .base import DocumentFormat, DocumentParser, RawDocument, RawParagraph

logger = logging.getLogger(__name__)

# Patrones para detectar capítulos
CHAPTER_PATTERNS = [
    re.compile(r"^cap[íi]tulo\s+\d+", re.IGNORECASE),
    re.compile(r"^chapter\s+\d+", re.IGNORECASE),
    re.compile(r"^parte\s+\d+", re.IGNORECASE),
    # Numeración romana: Debe ser un número romano válido (no solo cualquier letra)
    # y estar solo o seguido de un título corto (no una oración larga)
    re.compile(
        r"^(I{1,3}|IV|V|VI{0,3}|IX|X{1,3}|XI{0,3}|XIV|XV|XVI{0,3}|XIX|XX{0,3})\.\s*$", re.IGNORECASE
    ),
    re.compile(r"^\d+\.\s+[A-ZÁÉÍÓÚÑ]"),  # "1. Título"
]

# Patrón para detectar número de capítulo solo (sin título en la misma línea)
CHAPTER_NUMBER_ONLY_PATTERN = re.compile(r"^(\d{1,3})\s*$")

# Conectores narrativos que indican que un párrafo es contenido, no título
NARRATIVE_STARTERS = [
    "cuando",
    "mientras",
    "entonces",
    "sin embargo",
    "no obstante",
    "de pronto",
    "aquella",
    "aquel",
    "aquellos",
    "aquellas",
    "había",
    "era",
    "fue",
    "estaba",
    "tenía",
    "hacía",
    "el día",
    "la noche",
    "esa mañana",
    "aquella tarde",
    "después de",
    "antes de",
    "al cabo de",
    "tras",
    "a la mañana",
    "por la noche",
    "durante",
    "desde que",
]

# Longitud máxima para considerar un texto como posible título
MAX_TITLE_LENGTH = 60


def is_likely_chapter_title(text: str) -> bool:
    """
    Determina si un texto es probablemente un título de capítulo.

    Heurísticas:
    1. Longitud corta (<60 caracteres)
    2. No empieza con conectores narrativos
    3. No contiene verbos conjugados comunes al inicio

    Args:
        text: Texto a evaluar

    Returns:
        True si parece ser un título
    """
    text = text.strip()

    # Demasiado largo para título
    if len(text) > MAX_TITLE_LENGTH:
        return False

    # Demasiado corto (solo puntuación o similar)
    if len(text) < 2:
        return False

    text_lower = text.lower()

    # Verificar conectores narrativos
    return all(not text_lower.startswith(starter) for starter in NARRATIVE_STARTERS)


class DocxParser(DocumentParser):
    """
    Parser para documentos DOCX.

    Extrae:
    - Párrafos con posiciones de caracteres
    - Headings y niveles
    - Metadatos del documento
    - Detección de posibles capítulos
    """

    format = DocumentFormat.DOCX

    def __init__(self):
        self._docx_available = self._check_dependency()

    def _check_dependency(self) -> bool:
        """Verifica que python-docx está disponible."""
        try:
            import docx

            return True
        except ImportError:
            logger.warning("python-docx no instalado. Instalar con: pip install python-docx")
            return False

    def parse(self, path: Path) -> Result[RawDocument]:
        """
        Parsea un documento DOCX.

        Args:
            path: Ruta al archivo DOCX

        Returns:
            Result con RawDocument o error
        """
        # Validar archivo antes de abrir
        validation_result = self.validate_file(path)
        if validation_result.is_failure:
            return validation_result

        path = validation_result.value

        if not self._docx_available:
            from ..core.errors import ErrorSeverity, NarrativeError

            return Result.failure(
                NarrativeError(
                    message="python-docx not installed",
                    severity=ErrorSeverity.FATAL,
                    user_message="Instala python-docx: pip install python-docx",
                )
            )

        import docx

        try:
            doc = docx.Document(str(path))
        except BadZipFile:
            return Result.failure(
                CorruptedDocumentError(
                    file_path=str(path),
                    original_error="Archivo ZIP inválido (no es un DOCX válido)",
                )
            )
        except Exception as e:
            return Result.failure(
                CorruptedDocumentError(
                    file_path=str(path),
                    original_error=str(e),
                )
            )

        # Extraer metadatos
        metadata = self._extract_metadata(doc)

        # Extraer párrafos
        paragraphs = []
        # Las posiciones se recalcularán por TextCoordinateSystem en RawDocument
        # para garantizar consistencia con full_text

        # Verificar si hay comentarios en el documento
        has_comments = self._has_comments(doc)
        if has_comments:
            logger.info("Documento contiene comentarios/anotaciones (serán ignorados)")

        # Verificar si hay Track Changes
        has_track_changes = self._has_track_changes(doc)
        if has_track_changes:
            logger.info(
                "Documento contiene revisiones (Track Changes). Se aceptarán automáticamente."
            )

        for idx, para in enumerate(doc.paragraphs):
            # Obtener texto LIMPIO (sin tachados, aceptando Track Changes)
            clean_text = self._get_clean_text(para)

            # Detectar si es heading
            is_heading, heading_level = self._detect_heading(para)

            raw_para = RawParagraph(
                text=clean_text,  # Texto limpio
                index=idx,
                style_name=para.style.name if para.style else "Normal",
                is_heading=is_heading,
                heading_level=heading_level,
                # Posiciones temporales - serán recalculadas por TextCoordinateSystem
                start_char=0,
                end_char=len(clean_text),
                metadata={
                    "has_bold": self._has_bold(para),
                    "has_italic": self._has_italic(para),
                    "alignment": str(para.alignment) if para.alignment else None,
                    "has_strikethrough": self._has_strikethrough(para),
                    "has_track_changes": self._paragraph_has_track_changes(para),
                    "original_text": para.text if clean_text != para.text else None,
                },
            )

            paragraphs.append(raw_para)

        # Post-proceso: detectar títulos de capítulo cuando hay número solo en una línea
        self._detect_chapter_titles(paragraphs)

        # Verificar que hay contenido
        total_text = sum(len(p.text.strip()) for p in paragraphs)
        if total_text == 0:
            return Result.failure(EmptyDocumentError(file_path=str(path)))

        # Crear documento
        raw_doc = RawDocument(
            paragraphs=paragraphs,
            metadata=metadata,
            source_path=path,
            source_format=DocumentFormat.DOCX,
        )

        # Advertencias
        if self._has_tables(doc):
            raw_doc.add_warning(
                "El documento contiene tablas. El contenido de las tablas puede no extraerse correctamente."
            )

        if self._has_images(doc):
            raw_doc.add_warning("El documento contiene imágenes que serán ignoradas.")

        logger.info(
            f"DOCX parseado: {raw_doc.paragraph_count} párrafos, "
            f"{raw_doc.word_count} palabras, {len(raw_doc.headings)} headings"
        )

        return Result.success(raw_doc)

    def _extract_metadata(self, doc) -> dict:
        """Extrae metadatos del documento."""
        props = doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "last_modified_by": props.last_modified_by or "",
        }

    def _detect_heading(self, para) -> tuple[bool, int | None]:
        """
        Detecta si un párrafo es un heading.

        Returns:
            Tupla (is_heading, level)
        """
        style_name = para.style.name if para.style else ""

        # Por estilo
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
                return True, level
            except ValueError:
                return True, 1

        # Por patrón de texto (capítulos)
        text = para.text.strip()
        for pattern in CHAPTER_PATTERNS:
            if pattern.match(text):
                return True, 1

        return False, None

    def _get_clean_text(self, para) -> str:
        """
        Obtiene el texto del párrafo procesando Track Changes y strikethrough.

        Maneja:
        - Texto tachado (strikethrough): excluido
        - Track Changes (revisiones de Word):
          - w:del (texto eliminado): excluido
          - w:ins (texto insertado): incluido (se acepta la revisión)

        Args:
            para: Párrafo de python-docx

        Returns:
            Texto limpio aceptando todas las revisiones
        """
        # Primero intentar con Track Changes via XML
        clean_text = self._get_text_with_track_changes(para)
        if clean_text is not None:
            return clean_text

        # Fallback: método tradicional con strikethrough
        clean_parts = []

        for run in para.runs:
            # Verificar si el run tiene strikethrough
            is_struck = False
            try:
                if run.font.strike:
                    is_struck = True
            except Exception:
                pass

            # También verificar double strikethrough
            try:
                if run.font.double_strike:
                    is_struck = True
            except Exception:
                pass

            if not is_struck:
                clean_parts.append(run.text)

        return "".join(clean_parts)

    def _get_text_with_track_changes(self, para) -> str | None:
        """
        Extrae texto del párrafo procesando Track Changes via XML.

        Track Changes en DOCX:
        - <w:del>: Texto eliminado (debe excluirse)
        - <w:ins>: Texto insertado (debe incluirse)
        - <w:t>: Texto normal

        Args:
            para: Párrafo de python-docx

        Returns:
            Texto con revisiones aceptadas, o None si no hay Track Changes
        """
        try:
            from docx.oxml.ns import qn

            # Acceder al XML del párrafo
            p_element = para._element

            # Verificar si hay elementos de revisión
            has_revisions = p_element.findall(".//" + qn("w:del")) or p_element.findall(
                ".//" + qn("w:ins")
            )

            if not has_revisions:
                return None

            # Procesar el XML para extraer texto
            text_parts: list = []
            self._extract_text_from_element(p_element, text_parts, qn)

            return "".join(text_parts)

        except Exception as e:
            logger.debug(f"Error procesando Track Changes: {e}")
            return None

    def _extract_text_from_element(self, element, text_parts: list, qn) -> None:
        """
        Extrae texto recursivamente de un elemento XML, respetando Track Changes.

        Args:
            element: Elemento XML
            text_parts: Lista donde acumular partes de texto
            qn: Función qn de docx.oxml.ns
        """
        # Tags a ignorar completamente (texto eliminado)
        skip_tags = {qn("w:del"), qn("w:delText")}

        for child in element:
            tag = child.tag

            # Ignorar texto eliminado
            if tag in skip_tags:
                continue

            # Texto normal o insertado
            if tag == qn("w:t"):
                if child.text:
                    text_parts.append(child.text)

            # Texto insertado (w:ins) - procesar recursivamente
            elif tag == qn("w:ins"):
                self._extract_text_from_element(child, text_parts, qn)

            # Run (w:r) - procesar recursivamente
            elif tag == qn("w:r"):
                # Verificar si el run tiene strikethrough
                rpr = child.find(qn("w:rPr"))
                is_struck = False
                if rpr is not None:
                    strike = rpr.find(qn("w:strike"))
                    dstrike = rpr.find(qn("w:dstrike"))
                    if strike is not None or dstrike is not None:
                        is_struck = True

                if not is_struck:
                    self._extract_text_from_element(child, text_parts, qn)

            # Otros elementos - procesar recursivamente
            else:
                self._extract_text_from_element(child, text_parts, qn)

    def _has_strikethrough(self, para) -> bool:
        """Verifica si el párrafo tiene texto tachado."""
        for run in para.runs:
            try:
                if run.font.strike or run.font.double_strike:
                    return True
            except (AttributeError, TypeError):
                pass
        return False

    def _has_comments(self, doc) -> bool:
        """
        Verifica si el documento tiene comentarios/anotaciones.

        Los comentarios están en una parte separada del documento
        y no se incluyen en el texto de los párrafos normalmente.
        """
        try:
            # Los comentarios están en el namespace w:comments
            # Accedemos al XML subyacente

            # Buscar la parte de comentarios
            document_part = doc.part
            for rel in document_part.rels.values():
                if "comments" in str(rel.reltype).lower():
                    return True
        except Exception:
            pass
        return False

    def _has_bold(self, para) -> bool:
        """Verifica si el párrafo tiene texto en negrita."""
        return any(run.bold for run in para.runs)

    def _has_italic(self, para) -> bool:
        """Verifica si el párrafo tiene texto en cursiva."""
        return any(run.italic for run in para.runs)

    def _has_tables(self, doc) -> bool:
        """Verifica si el documento tiene tablas."""
        return len(doc.tables) > 0

    def _has_images(self, doc) -> bool:
        """Verifica si el documento tiene imágenes."""
        return any("image" in rel.reltype for rel in doc.part.rels.values())

    def _has_track_changes(self, doc) -> bool:
        """
        Verifica si el documento tiene Track Changes (revisiones).
        """
        try:
            from docx.oxml.ns import qn

            # Buscar elementos w:del o w:ins en el body
            body = doc.element.body
            has_del = body.findall(".//" + qn("w:del"))
            has_ins = body.findall(".//" + qn("w:ins"))
            return bool(has_del or has_ins)
        except Exception:
            return False

    def _paragraph_has_track_changes(self, para) -> bool:
        """
        Verifica si un párrafo específico tiene Track Changes.
        """
        try:
            from docx.oxml.ns import qn

            p_element = para._element
            has_del = p_element.findall(".//" + qn("w:del"))
            has_ins = p_element.findall(".//" + qn("w:ins"))
            return bool(has_del or has_ins)
        except Exception:
            return False

    def _detect_chapter_titles(self, paragraphs: list[RawParagraph]) -> None:
        """
        Post-proceso para detectar títulos de capítulo.

        Cuando hay un número solo en una línea, el siguiente párrafo
        podría ser el título del capítulo (si cumple las heurísticas).

        Modifica los párrafos in-place para marcar headings.
        """
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            text = para.text.strip()

            # Verificar si es un número de capítulo solo
            match = CHAPTER_NUMBER_ONLY_PATTERN.match(text)
            if match:
                # Marcar este párrafo como heading (número de capítulo)
                para.is_heading = True
                para.heading_level = 1
                para.metadata["chapter_number"] = int(match.group(1))

                # Buscar el siguiente párrafo no vacío
                next_idx = i + 1
                while next_idx < len(paragraphs) and paragraphs[next_idx].is_empty:
                    next_idx += 1

                if next_idx < len(paragraphs):
                    next_para = paragraphs[next_idx]
                    next_text = next_para.text.strip()

                    # Verificar si parece un título
                    if is_likely_chapter_title(next_text):
                        # Este párrafo es el título del capítulo, pero NO lo marcamos
                        # como heading para evitar duplicados en la detección de estructura.
                        # Solo añadimos metadata para identificarlo como título.
                        next_para.metadata["is_chapter_title"] = True
                        next_para.metadata["chapter_number"] = int(match.group(1))

                        # Combinar número y título en metadata del número (heading principal)
                        para.metadata["chapter_title"] = next_text

                        logger.debug(
                            f"Detectado capítulo {match.group(1)}: '{next_text[:50]}...'"
                            if len(next_text) > 50
                            else f"Detectado capítulo {match.group(1)}: '{next_text}'"
                        )

            i += 1
